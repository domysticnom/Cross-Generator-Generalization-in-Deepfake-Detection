"""
Bundle the repo's developer docs (docs/*.md) into ONE professional Word document,
styled to match Group14_Final_Proposal.docx.

The docs/*.md files stay in the repo as the working source (code links to them);
this just produces a single submittable Group14_Project_Documentation.docx by
converting and concatenating them, each as its own section.

Output:
    Group14_Project_Documentation.docx
"""

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# docs bundled in reading order (README.md is just an index -> skipped)
DOC_FILES = [
    ("Quickstart", "docs/QUICKSTART.md"),
    ("Dataset Access", "docs/DATASET_ACCESS.md"),
    ("Interfaces and Contracts", "docs/INTERFACES.md"),
    ("Preprocessing Cache Format", "docs/CACHE_FORMAT.md"),
]


# ---------------------------------------------------------------------------
# minimal markdown -> block parser (headings, lists, tables, code, quotes)
# ---------------------------------------------------------------------------
def parse_markdown(text):
    blocks = []
    lines = text.split("\n")
    i = 0
    para = []

    def flush_para():
        if para:
            blocks.append(("p", " ".join(para)))
            para.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            flush_para()
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(code)))
            i += 1
            continue

        # table (consecutive lines starting with |)
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_para()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                # skip the |---|---| separator row
                if not all(set(c) <= set("-: ") and c for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append(("table", rows))
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush_para()
            level = min(len(m.group(1)), 3)
            blocks.append((f"h{level}", m.group(2)))
            i += 1
            continue

        # horizontal rule
        if stripped in ("---", "***", "___"):
            flush_para()
            blocks.append(("hr", None))
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            flush_para()
            blocks.append(("quote", stripped.lstrip("> ").rstrip()))
            i += 1
            continue

        # bullet list
        if re.match(r"^[-*]\s+", stripped):
            flush_para()
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("bul", items))
            continue

        # numbered list
        if re.match(r"^\d+\.\s+", stripped):
            flush_para()
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("num", items))
            continue

        # blank line ends a paragraph
        if not stripped:
            flush_para()
            i += 1
            continue

        para.append(stripped)
        i += 1

    flush_para()
    return blocks


# ---------------------------------------------------------------------------
# inline rendering: **bold** and `code`
# ---------------------------------------------------------------------------
def add_inline(paragraph, text):
    for part in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def render_blocks(doc, blocks):
    for kind, payload in blocks:
        if kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "h3":
            doc.add_heading(payload, level=3)
        elif kind == "p":
            add_inline(doc.add_paragraph(), payload)
        elif kind == "bul":
            for it in payload:
                add_inline(doc.add_paragraph(style="List Bullet"), it)
        elif kind == "num":
            for it in payload:
                add_inline(doc.add_paragraph(style="List Number"), it)
        elif kind == "quote":
            p = doc.add_paragraph(style="Intense Quote") if "Intense Quote" in [s.name for s in doc.styles] else doc.add_paragraph()
            add_inline(p, payload)
        elif kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(payload)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif kind == "table":
            header, *rows = payload
            table = doc.add_table(rows=1, cols=len(header))
            try:
                table.style = "Light Grid Accent 1"
            except Exception:
                table.style = "Table Grid"
            for c, txt in zip(table.rows[0].cells, header):
                c.text = ""
                r = c.paragraphs[0].add_run(txt)
                r.font.bold = True
                r.font.size = Pt(9)
            for row in rows:
                cells = table.add_row().cells
                for c, txt in zip(cells, row):
                    c.text = ""
                    add_inline(c.paragraphs[0], txt)
                    for run in c.paragraphs[0].runs:
                        run.font.size = Pt(9)
        elif kind == "hr":
            doc.add_paragraph()


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # title page
    h = doc.add_heading("Project Documentation", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for sub in ("Cross-Generator Generalization in Deepfake Detection",
                "IE7374: Generative AI - Northeastern University - Group 14"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sub)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    for title, path in DOC_FILES:
        doc.add_page_break()
        with open(path, encoding="utf-8") as f:
            blocks = parse_markdown(f.read())
        # drop the file's own leading H1 (we use the section title below instead)
        if blocks and blocks[0][0] == "h1":
            blocks = blocks[1:]
        doc.add_heading(title, level=1)
        render_blocks(doc, blocks)

    out = "Group14_Project_Documentation.docx"
    doc.save(out)
    print("wrote", out, "from", len(DOC_FILES), "docs")


if __name__ == "__main__":
    main()
