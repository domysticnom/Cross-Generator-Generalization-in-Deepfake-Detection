"""
Generate the "Research and Selection of Methods" deliverable as a DOCX
(styled to match Group14_Final_Proposal.docx).

Single source of truth: the `BLOCKS` list below. Content is grounded in the
research written to .planning/research/{BACKGROUND,DATASETS,METHODOLOGY_CNN}.md
and the real preliminary experiments in the repo (preprocessing detection rates
on FF++ c23, and the hermetic tests under tests/). Every number traces to a
cited source or to an actual run/test.

Output:
    Group14_Research_and_Methods.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Content blocks. Each block is (kind, payload).
#   title/subtitle/h1/h2 -> headings
#   p            -> paragraph (supports **bold**)
#   bul / num    -> list of items (support **bold**)
#   table        -> [header_row, *data_rows]
#   refs         -> numbered reference list
#   spacer       -> blank line
# ---------------------------------------------------------------------------

BLOCKS = [
    ("title", "Cross-Generator Generalization in Deepfake Detection"),
    ("subtitle", "IE7374: Generative AI — Northeastern University · Group 14"),
    ("subtitle", "Research and Selection of Methods"),
    ("subtitle", "Dominic Rivas · Jonathan Jude Regalado · Lyxelis Rodriguez Navarro · "
                 "Obinna Okonkwo · Sagar Ayare"),
    ("spacer", None),

    ("p",
     "This document covers the research phase of the project: what we are trying to do, what the "
     "literature says, how we compared the candidate methods and datasets, and the small-scale "
     "experiments we ran to confirm the approach is feasible before committing to full training. "
     "The project's core claim is methodological: **cross-generator generalization must be "
     "measured, not assumed.** A detector that scores 99% on its own training distribution and 60% "
     "on an unseen generator is a more useful result than a single in-distribution number, and "
     "every choice below is made to support that measurement."),

    # ---------------------------------------------------------------- 1
    ("h1", "1. Define Objectives"),
    ("p",
     "The task is binary classification of face imagery as **real vs. manipulated**, at both the "
     "frame level (the model's native unit) and the video level (mean of frame scores). The "
     "project is not merely to train a deepfake detector — it is to **characterize how a "
     "conventional CNN detector transfers to manipulation methods it was never trained on.**"),
    ("p", "Concrete, measurable objectives:"),
    ("num", [
        "Train conventional CNN detectors — **EfficientNet** and **XceptionNet** — on "
        "FaceForensics++ (FF++), the canonical four-method benchmark.",
        "Run a rotating **leave-one-manipulation-out (LOMO)** protocol: train on 3 of "
        "{DeepFakes, Face2Face, FaceSwap, NeuralTextures}, test on the held-out 4th, rotating the "
        "held-out method to fill a **transfer matrix**.",
        "Report **seen vs. unseen separately** — never folded into one accuracy — with the "
        "**generalization gap (seen AUC − unseen AUC)** as the headline number.",
        "Enforce **identity-disjoint splits** so the unseen-method result is a clean measurement "
        "and not identity leakage.",
        "(Scope-gated stretch) Extend to a cross-dataset test (FF++ → DFDC) and a self-generated "
        "**SimSwap** unseen-generator set to see whether the pattern holds beyond FF++.",
    ]),
    ("p",
     "Why this task: deepfake detectors are almost always reported on a single in-distribution "
     "number, which hides the fact that they overfit to the specific generators in their training "
     "set. The interesting, under-reported quantity is what happens on an unseen generator — "
     "exactly what the transfer matrix measures. This is a directly generative-AI-framed question: "
     "each new generator family (GAN → diffusion, one manipulation method → another) is a "
     "distribution shift the detector was never trained on."),

    # ---------------------------------------------------------------- 2
    ("h1", "2. Literature Review"),
    ("p",
     "Every citation below was checked against a fetched primary or near-primary source. Numbers "
     "from a single paper or secondary summary are flagged for verification before the final "
     "report."),

    ("h2", "2.1 Foundational detectors (the baselines we build on)"),
    ("bul", [
        "**FaceForensics++ (Rössler et al., ICCV 2019).** Introduced the FF++ benchmark "
        "(1,000 pristine sequences × 4 manipulation methods) and established **XceptionNet** as "
        "the strong supervised baseline — 99.26% (raw), **95.73% (c23)**, 81.00% (c40) accuracy.",
        "**MesoNet (Afchar et al., WIFS 2018).** A deliberately shallow CNN targeting mesoscopic "
        "artifacts; the lightweight predecessor to deeper-CNN detectors.",
        "**EfficientNet (Tan & Le, ICML 2019).** Compound model scaling; became the dominant "
        "backbone in deepfake-detection competitions.",
        "**DFDC winning solution (Seferbekov, 2020).** MTCNN faces + an ensemble of seven "
        "EfficientNet-B7 models; scored only **~65% accuracy on the held-out private set** despite "
        "topping the public set — a concrete demonstration that leaderboard performance does not "
        "transfer.",
        "**Face X-Ray (Li et al., CVPR 2020).** Detects the blending boundary common to face swaps "
        "rather than method-specific artifacts, explicitly targeting generalization.",
        "**Frequency analysis (Frank et al., ICML 2020).** GAN images carry systematic "
        "frequency-domain artifacts from upsampling; context for why GAN-tuned detectors fail on "
        "diffusion.",
        "**Capsule-Forensics (Nguyen et al., ICASSP 2019).** Capsule networks for forgery "
        "detection; the architecture-diversity strand of early detectors.",
    ]),

    ("h2", "2.2 The generalization problem (the gap we measure)"),
    ("p",
     "These works show, with numbers, that strong in-distribution detectors collapse on unseen "
     "generators or datasets:"),
    ("table", [
        ["Setting", "Detector", "In-distribution", "Unseen", "Source"],
        ["FF++ → Celeb-DF (cross-dataset)", "Xception", "99.7 AUC", "48.2 AUC",
         "Li et al., Celeb-DF 2020"],
        ["FF++ leave-one-out (cross-manip)", "Xception", "high on seen",
         "77.9% avg; 51.2% on FaceSwap", "Haliassos et al., 2021"],
        ["FF++ leave-one-out (same protocol)", "LipForensics", "—", "97.1% avg",
         "Haliassos et al., 2021"],
        ["FF++(+BI) → unseen datasets", "Face X-Ray", "high on FF++",
         "80.9 / 80.6 / 95.4 AUC", "Li et al., Face X-Ray 2020"],
        ["GAN → diffusion (cross-family)", "GAN-trained CNN", "~97.4% AUROC",
         "~78.6% AUROC; 26% recall", "Ricker et al., 2024"],
        ["DFDC public → private", "EfficientNet-B7", "top public score",
         "~65% acc", "Seferbekov, 2020"],
    ]),
    ("p",
     "**Pattern:** every line has the same shape — strong (95–99%) in-distribution performance "
     "degrading to roughly 48–80% on an unseen generator, often approaching chance (50% AUC) for "
     "the hardest cases. The size of the drop is the result of interest, not the in-distribution "
     "peak."),
    ("p",
     "The cleanest single piece of evidence is **LipForensics (CVPR 2021)**: under the exact LOMO "
     "protocol we adopt, the Xception baseline it reports averages **77.9% AUC** and falls to "
     "**51.2% AUC (near chance) on the unseen FaceSwap** manipulation — while a "
     "generalization-oriented design reaches 97.1%. That gap is what we set out to reproduce and "
     "characterize with our own detectors."),

    ("h2", "2.3 Gap statement"),
    ("p",
     "Standard CNN detectors such as XceptionNet and EfficientNet reach near-saturated accuracy "
     "when trained and tested on the same manipulation methods, yet collapse toward chance on "
     "generators they were never trained on (FF++ Xception 99.7 → 48.2 AUC on Celeb-DF; → 51.2 AUC "
     "on unseen FaceSwap; GAN-trained detectors retain only ~26% recall on diffusion faces). "
     "Despite this, detectors are still most often reported on a single in-distribution number. "
     "**This project treats cross-generator generalization as a quantity to be measured rather "
     "than assumed** — training conventional CNN detectors on FF++ and explicitly reporting how "
     "their performance transfers to manipulation methods held out of training."),

    # ---------------------------------------------------------------- 3
    ("h1", "3. Benchmarking and Selection"),
    ("p",
     "We compared candidate models, datasets, and the evaluation protocol on the criteria that "
     "matter for a one-semester, single-GPU project."),

    ("h2", "3.1 Model selection"),
    ("p",
     "We commit to two ImageNet-pretrained, single-frame CNN classifiers — exactly the "
     "\"conventional CNN detector\" class the project characterizes. Comparing the two also adds a "
     "within-project axis: the architecture-sensitivity of generalization."),
    ("table", [
        ["Criterion", "XceptionNet", "EfficientNet (B0–B4)"],
        ["In-distribution accuracy", "0.9637 AUC (FF++ c23)", "0.9567 AUC (B4, same benchmark)"],
        ["Provenance", "The FF++ reference detector", "The DFDC-winning family"],
        ["Computational efficiency", "~22.9M params, native 299×299",
         "B0 ~5.3M / B4 ~19M params, compound-scaled"],
        ["Scalability", "Fixed architecture", "Family B0→B7 trades accuracy for compute"],
        ["Pretrained availability", "ImageNet via timm", "ImageNet (+ Noisy Student) via timm"],
    ]),
    ("p",
     "**Rejected alternatives:** MesoNet (too shallow), Capsule-Forensics (niche), and the "
     "generalization-*specific* designs Face X-Ray / LipForensics. The last two are deliberately "
     "excluded from the detector choice — they are engineered to close the generalization gap, "
     "whereas our contribution is to **measure** the gap for a plain CNN. They belong in the "
     "discussion as the \"what a generalization-oriented method would recover\" contrast, not as "
     "our model. The realistic in-distribution target to report is DeepfakeBench's **~0.96 AUC at "
     "c23**, not the FF++ paper's headline ~0.99 (easier protocol, different splits)."),

    ("h2", "3.2 Dataset selection"),
    ("table", [
        ["Criterion", "FaceForensics++ (primary)", "DFDC Preview (stretch)"],
        ["Fit to the question", "4 clean methods → ideal LOMO transfer matrix",
         "Undisclosed generators disjoint from FF++"],
        ["Structure", "1,000 reals × 4 methods = 5,000 videos; identity-disjoint split",
         "~5,000 clips, 66 actors, 2 swap methods"],
        ["Size / feasibility", "c23 ~20 GB combined — tractable",
         "Full DFDC ~470 GB (excluded); Preview ~5K usable"],
        ["Binding constraint", "Face-cropping time, not GPU", "Download + crop time"],
        ["Access", "Google form (research license); c23 Kaggle mirror",
         "Kaggle competition rules"],
    ]),
    ("p",
     "FF++ is the primary training corpus because its **four generators are exactly the axis the "
     "transfer matrix rotates over**, and its official split is identity-disjoint by construction. "
     "DFDC is scope-gated because data-prep time — not compute — is the semester's binding "
     "constraint. A self-generated **SimSwap** set is a third, optional unseen-generator column "
     "(the evaluator skips it gracefully until it exists). A useful structural property of FF++: "
     "the four methods form a clean **2×2 design** — SWAP = {DeepFakes (learned), FaceSwap "
     "(graphics)}, REENACTMENT = {NeuralTextures (learned), Face2Face (graphics)} — letting the "
     "transfer matrix speak to both swap-vs-reenactment and learned-vs-graphics transfer."),

    ("h2", "3.3 Protocol, metrics, and expected results"),
    ("bul", [
        "**Protocol:** rotating LOMO (train-on-3 / test-on-1). We use train-on-3, which generally "
        "*narrows* the gap versus the train-on-1 rows common in the literature — a distinction to "
        "state so reviewers do not compare apples to oranges.",
        "**Primary metric:** AUC/AUROC (threshold-free, robust to the 1:4 real:fake imbalance), at "
        "both frame and video level.",
        "**Secondary:** accuracy, precision, recall, F1, per-split confusion matrices.",
        "**Fixed compression:** report at c23 (HQ) throughout.",
        "**Headline artifact:** a 4×(held-out) transfer matrix of AUCs with the in-distribution "
        "diagonal, a ΔAUC (seen − unseen) column, and ROC curves overlaying seen vs. unseen.",
    ]),
    ("p",
     "**Expected outcome (from the literature, to be confirmed by our runs):** in-distribution "
     "~0.95–0.99 AUC collapsing to roughly **~0.55–0.85 on unseen methods** — a 15–40 AUC-point "
     "drop, with NeuralTextures expected hardest to detect when held out and FaceSwap the weakest "
     "training source for transfer."),

    ("h2", "3.4 Feasibility verdict"),
    ("p",
     "Clearly tractable on a single ~12 GB GPU in one semester: ~5,000 videos → ~80k–160k cached "
     "face crops; fine-tuning a pretrained CNN is a few epochs / a few hours per run; the full "
     "protocol is **4 held-out folds × 2 backbones = 8 runs**, embarrassingly parallel across a "
     "5-person team sharing one cached crop store. Preprocessing (face cropping), not GPU time, is "
     "the binding constraint — which is why crops are cached to disk once."),

    # ---------------------------------------------------------------- 4
    ("h1", "4. Preliminary Experiments"),
    ("p",
     "Before committing to the 8 full training runs, we validated the pipeline with small-scale "
     "tests. These are real, reproducible checks in the repo — not the final results (no training "
     "run has completed yet; the transfer matrix is pending)."),

    ("h2", "4.1 Face-detection / preprocessing validated on real FF++ c23"),
    ("p",
     "We ran the full preprocessing pass on the real FF++ c23 corpus (all 5,000 clips, 20 frames "
     "each). The face-detection success rate was high and uniform across methods, confirming the "
     "crop pipeline is sound before any training:"),
    ("table", [
        ["Split", "Mean detection rate"],
        ["real (label 0)", "0.9726"],
        ["DeepFakes", "0.9717"],
        ["Face2Face", "0.9728"],
        ["FaceSwap", "0.9728"],
        ["NeuralTextures", "0.9732"],
    ]),
    ("p",
     "~97% of sampled frames yield a usable face crop across every method (≈97k crops from 5,000 "
     "clips), with no method-specific detection bias that would confound the cross-generator "
     "comparison. This is our first real experimental result and confirms the data path "
     "end-to-end."),

    ("h2", "4.2 Automated pipeline tests (7 passing)"),
    ("p",
     "The pipeline is covered by hermetic unit tests that run with no GPU, no ffmpeg, no "
     "mediapipe, and no corpus (using injected synthetic frames and a fake detector), so any "
     "teammate can validate the logic in seconds."),
    ("p", "**tests/test_preprocess.py (5 tests):**"),
    ("num", [
        "**Schema and crops** — the manifest matches the committed column schema; real is label 0 "
        "and every manipulation label 1; source_id is the identity prefix; a referenced crop is a "
        "real (size, size, 3) .npy on disk.",
        "**Detection-log math** — per-clip detection_rate = faces_detected / frames_sampled; a "
        "forced all-miss clip records rate 0.0; real/fake means are summarized separately.",
        "**Resume adds no duplicates** — re-running with --resume skips done work and never "
        "duplicates a crop_id or ledger row.",
        "**Copy-and-consume portability** — crops produced by preprocessing load correctly through "
        "the training CropDataset, returning (3, size, size) tensors.",
        "**Same clip_id across two methods** — FF++ reuses source-pair filenames across methods, "
        "so the resume ledger is keyed on (method, clip_id); both methods' crops are cached rather "
        "than one wrongly skipping the other.",
    ]),
    ("p",
     "**tests/test_inventory_ffpp.py (2 tests):** the inventory/integrity check returns OK with "
     "the exit-code contract on a complete tree, and flags MISMATCH (non-zero exit) when a method "
     "is short a clip. Together these lock down the two subtle correctness risks found during "
     "development — the (method, clip_id) resume key and identity-disjoint labeling."),

    ("h2", "4.3 Aggregation and environment checks"),
    ("bul", [
        "The transfer-matrix aggregator and assembly notebook were exercised on **synthetic "
        "per-run result JSONs**, confirming the end-of-project figure assembles correctly before "
        "any real run exists.",
        "**check_env.py** smoke-tests Python, torch/CUDA, ffmpeg, and core packages before any "
        "compute is spent, so a new machine fails fast and legibly.",
        "The crop cache uses a portable format with deterministic configs (seed 1337), so runs are "
        "reproducible across the team's heterogeneous hardware.",
    ]),

    ("h2", "4.4 Adjustments made from these experiments"),
    ("bul", [
        "The resume ledger was re-keyed from clip_id to **(method, clip_id)** after a test showed "
        "a clip_id-only key under-caches the fakes (same filename appears under multiple methods).",
        "Preprocessing defaults (20 frames, 256 px, seed 1337) are treated as config-driven "
        "placeholders — cheap to revise if an EDA pass changes them — rather than hard-coded.",
        "Uniform ~97% detection rates across methods confirmed we do not need a method-specific "
        "detector or per-method frame budget, simplifying the pipeline.",
    ]),

    # ---------------------------------------------------------------- 5
    ("h1", "5. Status and Open Items"),
    ("p", "**Runs complete:** 0 of 8 (preprocessing validated; training pending on team GPUs)."),
    ("p", "To verify before the final report:"),
    ("num", [
        "Pin the originating table for the Xception cross-manipulation numbers (they appear "
        "consistently across generalization papers; cite the primary source).",
        "Confirm we headline DeepfakeBench's ~0.96 in-distribution AUC (not the FF++ paper's "
        "~0.99).",
        "State the train-on-3 vs. train-on-1 distinction explicitly, since our LOMO protocol is "
        "expected to narrow the gap relative to the single-source rows quoted from the literature.",
    ]),

    # ---------------------------------------------------------------- refs
    ("h1", "References"),
    ("refs", [
        "Afchar, D., Nozick, V., Yamagishi, J., & Echizen, I. (2018). MesoNet: A compact facial "
        "video forgery detection network. IEEE WIFS. https://arxiv.org/abs/1809.00888",
        "Chollet, F. (2017). Xception: Deep learning with depthwise separable convolutions. CVPR. "
        "https://arxiv.org/abs/1610.02357",
        "Dolhansky, B., et al. (2019). The Deepfake Detection Challenge (DFDC) Preview dataset. "
        "https://arxiv.org/abs/1910.08854",
        "Dolhansky, B., et al. (2020). The DeepFake Detection Challenge (DFDC) dataset. "
        "https://arxiv.org/abs/2006.07397",
        "Frank, J., et al. (2020). Leveraging frequency analysis for deep fake image recognition. "
        "ICML. https://arxiv.org/abs/2003.08685",
        "Haliassos, A., Vougioukas, K., Petridis, S., & Pantic, M. (2021). Lips don't lie: A "
        "generalisable and robust approach to face forgery detection (LipForensics). CVPR. "
        "https://arxiv.org/abs/2012.07657",
        "Li, L., et al. (2020). Face X-Ray for more general face forgery detection. CVPR. "
        "https://arxiv.org/abs/1912.13458",
        "Li, Y., Yang, X., Sun, P., Qi, H., & Lyu, S. (2020). Celeb-DF: A large-scale challenging "
        "dataset for DeepFake forensics. CVPR. https://arxiv.org/abs/1909.12962",
        "Nguyen, H. H., Yamagishi, J., & Echizen, I. (2019). Capsule-Forensics. ICASSP. "
        "https://arxiv.org/abs/1810.11215",
        "Ricker, J., Damm, S., Holz, T., & Fischer, A. (2024). Towards the detection of diffusion "
        "model deepfakes. VISAPP. https://arxiv.org/abs/2210.14571",
        "Rössler, A., et al. (2019). FaceForensics++: Learning to detect manipulated facial "
        "images. ICCV. https://arxiv.org/abs/1901.08971",
        "Seferbekov, S. (2020). DFDC winning solution (1st place). "
        "https://github.com/selimsef/dfdc_deepfake_challenge",
        "Tan, M., & Le, Q. V. (2019). EfficientNet: Rethinking model scaling for convolutional "
        "neural networks. ICML. https://arxiv.org/abs/1905.11946",
        "Yan, Z., et al. (2023). DeepfakeBench: A comprehensive benchmark of deepfake detection. "
        "NeurIPS Datasets & Benchmarks. https://arxiv.org/abs/2307.01426",
    ]),
]


# ---------------------------------------------------------------------------
# DOCX rendering (with a tiny **bold** parser) — matches generate_enhanced_proposal.py
# ---------------------------------------------------------------------------
def add_runs(paragraph, text):
    for i, chunk in enumerate(text.split("**")):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if i % 2 == 1:
            run.font.bold = True


def render_docx(blocks, path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for kind, payload in blocks:
        if kind == "title":
            h = doc.add_heading(payload, level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "subtitle":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(payload)
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        elif kind == "h1":
            doc.add_heading(payload, level=1)
        elif kind == "h2":
            doc.add_heading(payload, level=2)
        elif kind == "p":
            add_runs(doc.add_paragraph(), payload)
        elif kind == "bul":
            for it in payload:
                add_runs(doc.add_paragraph(style="List Bullet"), it)
        elif kind == "num":
            for it in payload:
                add_runs(doc.add_paragraph(style="List Number"), it)
        elif kind == "refs":
            for it in payload:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                p.paragraph_format.first_line_indent = Pt(-18)
                add_runs(p, it)
        elif kind == "table":
            header, *rows = payload
            table = doc.add_table(rows=1, cols=len(header))
            try:
                table.style = "Light Grid Accent 1"
            except Exception:
                table.style = "Table Grid"
            for c, text in zip(table.rows[0].cells, header):
                c.text = ""
                run = c.paragraphs[0].add_run(text)
                run.font.bold = True
                run.font.size = Pt(10)
            for r in rows:
                cells = table.add_row().cells
                for c, text in zip(cells, r):
                    c.text = ""
                    run = c.paragraphs[0].add_run(text)
                    run.font.size = Pt(10)
        elif kind == "spacer":
            doc.add_paragraph()

    doc.save(path)


if __name__ == "__main__":
    render_docx(BLOCKS, "Group14_Research_and_Methods.docx")
    print("Wrote Group14_Research_and_Methods.docx")
